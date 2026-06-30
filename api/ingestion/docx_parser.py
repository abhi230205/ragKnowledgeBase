"""DOCX text extraction (python-docx), feeding the same chunk/embed pipeline as PDFs.

Word documents have no fixed page model, so all text (paragraphs + table rows) is
returned as a single page (page 1). Citations therefore show p.1 for DOCX sources.
Corrupt/unreadable files raise ValueError so the caller records a per-file error
without aborting the batch — mirroring pdf_parser.
"""

from __future__ import annotations

import io

from ingestion.pdf_parser import PageText


def extract_pages(docx_bytes: bytes) -> list[PageText]:
    """Extract text from DOCX bytes as a single PageText (page 1)."""
    from docx import Document  # python-docx

    try:
        doc = Document(io.BytesIO(docx_bytes))
    except Exception as exc:
        raise ValueError(f"Could not open DOCX: {exc}") from exc

    parts: list[str] = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            # python-docx returns the SAME cell object for every grid position a
            # merged cell spans, so a horizontally merged cell would repeat its text
            # ("X | X"). De-dupe consecutive identical underlying cells (_tc).
            cells: list[str] = []
            prev_tc = None
            for c in row.cells:
                if c._tc is prev_tc:
                    continue
                prev_tc = c._tc
                cells.append(c.text.strip().replace("\n", " "))
            if any(cells):
                parts.append(" | ".join(cells))

    return [PageText(page_number=1, text="\n".join(parts).strip())]
