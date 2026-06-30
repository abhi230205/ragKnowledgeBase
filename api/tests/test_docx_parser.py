"""Unit tests for the DOCX parser (python-docx -> PageText).

Builds real .docx bytes in memory and checks paragraph + table extraction, the
single-page model (DOCX has no page breaks we can rely on), and that corrupt
input raises ValueError so the pipeline records a per-file error rather than
crashing the batch.
"""

from __future__ import annotations

import io

import pytest
from docx import Document

from ingestion import docx_parser


def _docx_bytes(paragraphs: list[str], table: list[list[str]] | None = None) -> bytes:
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if table:
        t = doc.add_table(rows=len(table), cols=len(table[0]))
        for i, row in enumerate(table):
            for j, cell in enumerate(row):
                t.cell(i, j).text = cell
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_pages_returns_single_page_with_text():
    pages = docx_parser.extract_pages(_docx_bytes(["Hello world", "Second paragraph."]))
    assert len(pages) == 1
    assert pages[0].page_number == 1
    assert "Hello world" in pages[0].text
    assert "Second paragraph." in pages[0].text


def test_extract_pages_includes_table_cells():
    pages = docx_parser.extract_pages(
        _docx_bytes(["Intro"], table=[["Plan", "Price"], ["Pro", "$99"]])
    )
    text = pages[0].text
    assert "Plan | Price" in text
    assert "Pro | $99" in text


def test_extract_pages_skips_blank_paragraphs():
    pages = docx_parser.extract_pages(_docx_bytes(["", "   ", "Real content"]))
    assert pages[0].text == "Real content"


def test_extract_pages_dedupes_merged_table_cells():
    """python-docx repeats the same cell object across a merge span; we de-dupe so a
    horizontally merged cell renders once ('SPAN | Z'), not 'SPAN | SPAN | Z'."""
    doc = Document()
    t = doc.add_table(rows=1, cols=3)
    row = t.rows[0]
    merged = row.cells[0].merge(row.cells[1])
    merged.text = "SPAN"
    row.cells[-1].text = "Z"
    buf = io.BytesIO()
    doc.save(buf)

    pages = docx_parser.extract_pages(buf.getvalue())
    assert "SPAN | Z" in pages[0].text
    assert "SPAN | SPAN" not in pages[0].text


def test_extract_pages_raises_on_corrupt_bytes():
    with pytest.raises(ValueError):
        docx_parser.extract_pages(b"this is not a real docx file")
